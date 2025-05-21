# app/models/report_generator.py
import os
import logging
import time
import json
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import io
import base64
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Union, Any
import uuid

from flask import current_app, url_for
from werkzeug.utils import secure_filename

# Import project modules
from app.models.analysis import AnalysisMetadata
from app.utilities import (
   convert_to_json_serializable, 
   get_full_variable_name, 
   get_reports_folder,
   safe_path_join,
   log_error_details,
   create_unique_filename,
   timeit
)
from app.ai_utils import get_llm_manager

# Set up logging
logger = logging.getLogger(__name__)

class ReportGenerator:
   """
   Enhanced report generator class with LLM-powered insights
   and rich metadata integration
   """
   
   def __init__(self, data_handler, metadata=None):
       """
       Initialize report generator with data handler and optional metadata
       
       Args:
           data_handler: DataHandler instance with analysis results
           metadata: Optional AnalysisMetadata instance for context
       """
       self.data_handler = data_handler
       self.metadata = metadata or self._get_default_metadata()
       self.llm_manager = get_llm_manager()
       self.session_id = os.path.basename(data_handler.session_folder)
       self.report_folder = get_reports_folder(self.session_id)
       
       # Format configuration
       self.format_config = {
           'color_primary': '#003366',
           'color_secondary': '#0077cc',
           'color_accent': '#00a8e8',
           'color_success': '#28a745',
           'color_warning': '#ffc107',
           'color_danger': '#dc3545',
           'color_text': '#333333',
           'color_text_light': '#666666',
           'font_main': 'Arial',
           'font_header': 'Arial',
           'logo_path': os.path.join(current_app.root_path, 'static', 'img', 'logo.png'),
           'max_width': 800,
           'page_margin': 20
       }
   
   def _get_default_metadata(self):
       """Get default metadata if none provided"""
       try:
           return AnalysisMetadata(
               session_id=os.path.basename(self.data_handler.session_folder)
           )
       except Exception as e:
           logger.warning(f"Could not create default metadata: {str(e)}")
           return None
   
   @timeit
   def generate_report(self, format_type='pdf', custom_sections=None, detail_level='standard'):
       """
       Generate a comprehensive report with LLM-enhanced insights
       
       Args:
           format_type: Report format (pdf, html, docx, md)
           custom_sections: Optional list of custom sections to include
           detail_level: Level of detail (basic, standard, technical)
           
       Returns:
           dict: Status and report information
       """
       try:
           # Validate that analysis has been run
           if not hasattr(self.data_handler, 'vulnerability_rankings') or self.data_handler.vulnerability_rankings is None:
               return {
                   'status': 'error',
                   'message': 'Vulnerability analysis has not been run. Run the analysis first.'
               }
           
           # Gather data for the report
           report_data = self._gather_report_data(detail_level)
           
           # Generate LLM insights for each section
           report_data = self._enhance_with_llm_insights(report_data, detail_level)
           
           # Add custom sections if provided
           if custom_sections:
               for section in custom_sections:
                   if section not in report_data['sections']:
                       report_data['sections'][section] = self._generate_custom_section(section)
           
           # Generate the report in the requested format
           if format_type.lower() == 'pdf':
               report_path, report_url = self._generate_pdf_report(report_data)
           elif format_type.lower() == 'html':
               report_path, report_url = self._generate_html_report(report_data)
           elif format_type.lower() == 'docx':
               report_path, report_url = self._generate_docx_report(report_data)
           elif format_type.lower() == 'md':
               report_path, report_url = self._generate_markdown_report(report_data)
           else:
               return {
                   'status': 'error',
                   'message': f'Unsupported report format: {format_type}'
               }
           
           # Log report generation
           self._log_report_generation(format_type, report_path)
           
           return {
               'status': 'success',
               'message': f'Successfully generated {format_type.upper()} report',
               'report_path': report_path,
               'report_url': report_url,
               'format': format_type,
               'timestamp': datetime.now().isoformat()
           }
           
       except Exception as e:
           error_info = log_error_details(e, session_id=self.session_id, 
                                        context={'report_type': format_type})
           return {
               'status': 'error',
               'message': f'Error generating report: {str(e)}',
               'error_details': error_info
           }
   
   def _gather_report_data(self, detail_level='standard'):
       """
       Gather all data needed for the report
       
       Args:
           detail_level: Level of detail (basic, standard, technical)
           
       Returns:
           dict: Data for the report
       """
       # Start with basic report data
       report_data = {
           'title': 'Malaria Risk Analysis Report',
           'subtitle': 'Spatial Vulnerability Assessment',
           'date': datetime.now().strftime('%B %d, %Y'),
           'timestamp': datetime.now().isoformat(),
           'detail_level': detail_level,
           'sections': {},
           'metadata': self._get_report_metadata()
       }
       
       # Add executive summary section
       report_data['sections']['executive_summary'] = self._generate_executive_summary()
       
       # Add data overview section
       report_data['sections']['data_overview'] = self._generate_data_overview()
       
       # Add methodology section with appropriate detail level
       report_data['sections']['methodology'] = self._generate_methodology_section(detail_level)
       
       # Add results section
       report_data['sections']['results'] = self._generate_results_section()
       
       # Add vulnerability rankings section
       report_data['sections']['vulnerability_rankings'] = self._generate_vulnerability_rankings()
       
       # Add urban extent analysis section if available
       if hasattr(self.data_handler, 'urban_extent_results') and self.data_handler.urban_extent_results:
           report_data['sections']['urban_extent'] = self._generate_urban_extent_section()
       
       # Add recommendations section
       report_data['sections']['recommendations'] = self._generate_recommendations()
       
       # Add appendix with technical details for higher detail levels
       if detail_level in ['standard', 'technical']:
           report_data['sections']['appendix'] = self._generate_appendix(detail_level)
       
       return report_data
   
   def _enhance_with_llm_insights(self, report_data, detail_level):
       """
       Enhance report data with LLM-generated insights
       
       Args:
           report_data: Base report data
           detail_level: Level of detail (basic, standard, technical)
           
       Returns:
           dict: Enhanced report data with LLM insights
       """
       # Only proceed if LLM manager is available
       if not self.llm_manager:
           logger.warning("LLM manager not available. Skipping insight generation.")
           return report_data
       
       try:
           # Create context for LLM insights
           llm_context = {
               'metadata': report_data['metadata'],
               'data_summary': {
                   'variable_count': len(self.data_handler.composite_variables or []),
                   'ward_count': len(self.data_handler.vulnerability_rankings) if hasattr(self.data_handler, 'vulnerability_rankings') else 0,
                   'analysis_type': 'custom' if hasattr(self.data_handler, 'custom_analysis') and self.data_handler.custom_analysis else 'standard'
               },
               'detail_level': detail_level
           }
           
           # Generate executive summary insights
           if 'executive_summary' in report_data['sections']:
               exec_summary_prompt = (
                   f"Generate an executive summary for a malaria risk analysis report. "
                   f"Detail level: {detail_level}. Include key findings, patterns identified, "
                   f"and main implications for intervention planning. "
                   f"The summary should be concise yet comprehensive."
               )
               insights = self.llm_manager.generate_response(
                   prompt=exec_summary_prompt, 
                   context={**llm_context, 'section': 'executive_summary'},
                   session_id=self.session_id
               )
               if insights:
                   report_data['sections']['executive_summary']['llm_insights'] = insights
           
           # Generate variable relationship insights
           if self.data_handler.composite_variables:
               variables_prompt = (
                   f"Analyze the relationship between the selected variables "
                   f"({', '.join(self.data_handler.composite_variables)}) "
                   f"and malaria risk. Identify potential interactions, synergies, "
                   f"or contradictions in how these variables affect vulnerability."
               )
               insights = self.llm_manager.generate_response(
                   prompt=variables_prompt, 
                   context={**llm_context, 'section': 'methodology'},
                   session_id=self.session_id
               )
               if insights:
                   if 'methodology' not in report_data['sections']:
                       report_data['sections']['methodology'] = {}
                   report_data['sections']['methodology']['variable_insights'] = insights
           
           # Generate vulnerability pattern insights
           if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
               # Get top and bottom 5 wards for context
               top_wards = self.data_handler.vulnerability_rankings.sort_values('overall_rank').head(5)['WardName'].tolist()
               bottom_wards = self.data_handler.vulnerability_rankings.sort_values('overall_rank', ascending=False).head(5)['WardName'].tolist()
               
               vuln_prompt = (
                   f"Analyze the vulnerability rankings. The most vulnerable wards are {', '.join(top_wards)} "
                   f"and the least vulnerable are {', '.join(bottom_wards)}. "
                   f"Identify spatial patterns, unexpected findings, and implications for intervention planning. "
                   f"Detail level: {detail_level}."
               )
               insights = self.llm_manager.generate_response(
                   prompt=vuln_prompt, 
                   context={**llm_context, 'section': 'vulnerability_rankings'},
                   session_id=self.session_id
               )
               if insights:
                   if 'vulnerability_rankings' not in report_data['sections']:
                       report_data['sections']['vulnerability_rankings'] = {}
                   report_data['sections']['vulnerability_rankings']['pattern_insights'] = insights
           
           # Generate urban extent insights if available
           if hasattr(self.data_handler, 'urban_extent_results') and self.data_handler.urban_extent_results:
               urban_insights = self._generate_urban_extent_insights(llm_context)
               if urban_insights:
                   if 'urban_extent' not in report_data['sections']:
                       report_data['sections']['urban_extent'] = {}
                   report_data['sections']['urban_extent']['llm_insights'] = urban_insights
           
           # Generate recommendation insights
           rec_prompt = (
               f"Based on the analysis results, generate key recommendations for "
               f"malaria intervention planning. Consider both geographical prioritization "
               f"and intervention strategy selection. Detail level: {detail_level}."
           )
           insights = self.llm_manager.generate_response(
               prompt=rec_prompt, 
               context={**llm_context, 'section': 'recommendations'},
               session_id=self.session_id
           )
           if insights:
               if 'recommendations' not in report_data['sections']:
                   report_data['sections']['recommendations'] = {}
               report_data['sections']['recommendations']['llm_insights'] = insights
           
           return report_data
       
       except Exception as e:
           logger.error(f"Error generating LLM insights: {str(e)}")
           # Return original data if there's an error
           return report_data
   
   def _generate_urban_extent_insights(self, llm_context):
       """
       Generate insights about urban extent analysis
       
       Args:
           llm_context: Context for LLM
           
       Returns:
           str: Insights about urban extent patterns
       """
       try:
           urban_thresholds = list(self.data_handler.urban_extent_results.keys())
           threshold_stats = {}
           
           for threshold in urban_thresholds:
               results = self.data_handler.urban_extent_results[threshold]
               urban_count = results['meets_threshold']
               rural_count = results['below_threshold']
               total = urban_count + rural_count
               urban_percent = (urban_count / total * 100) if total > 0 else 0
               
               threshold_stats[threshold] = {
                   'urban_count': urban_count,
                   'rural_count': rural_count,
                   'urban_percent': urban_percent
               }
           
           # Find non-urban high vulnerability wards
           non_urban_high_vuln = []
           if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
               # Get top 10 vulnerable wards
               top_vulnerable = self.data_handler.vulnerability_rankings.sort_values('overall_rank').head(10)
               
               # Default threshold for urban/rural classification
               default_threshold = 30
               
               # Check if these wards are urban at default threshold
               if default_threshold in self.data_handler.urban_extent_results:
                   results = self.data_handler.urban_extent_results[default_threshold]
                   for ward in top_vulnerable['WardName']:
                       if ward in results['below_threshold_wards']:
                           non_urban_high_vuln.append(ward)
           
           # Create prompt for urban insights
           urban_prompt = (
               f"Analyze the urban extent results across different thresholds: {urban_thresholds}. "
               f"At 30% threshold, {threshold_stats.get(30, {}).get('urban_percent', 0):.1f}% of wards are classified as urban. "
           )
           
           if non_urban_high_vuln:
               urban_prompt += (
                   f"The following high vulnerability wards are classified as non-urban: {', '.join(non_urban_high_vuln)}. "
                   f"What are the implications for intervention planning?"
               )
           
           # Generate insights
           insights = self.llm_manager.generate_response(
               prompt=urban_prompt, 
               context={
                   **llm_context, 
                   'section': 'urban_extent',
                   'threshold_stats': threshold_stats,
                   'non_urban_high_vuln': non_urban_high_vuln
               },
               session_id=self.session_id
           )
           
           return insights
       
       except Exception as e:
           logger.error(f"Error generating urban extent insights: {str(e)}")
           return None
   
   def _get_report_metadata(self):
       """Get metadata for the report"""
       metadata = {
           'generated_at': datetime.now().isoformat(),
           'session_id': self.session_id,
           'analysis_complete': True,
           'variables_used': self.data_handler.composite_variables if hasattr(self.data_handler, 'composite_variables') else [],
           'variable_count': len(self.data_handler.composite_variables) if hasattr(self.data_handler, 'composite_variables') else 0,
           'ward_count': len(self.data_handler.vulnerability_rankings) if hasattr(self.data_handler, 'vulnerability_rankings') else 0
       }
       
       # Add analysis steps if available
       if self.metadata and hasattr(self.metadata, 'steps'):
           metadata['analysis_steps'] = self.metadata.steps
       
       return metadata
   
   def _generate_executive_summary(self):
       """Generate executive summary section"""
       summary = {
           'title': 'Executive Summary',
           'content': []
       }
       
       # Add summary of the analysis
       if hasattr(self.data_handler, 'composite_variables') and self.data_handler.composite_variables:
           variables_used = self.data_handler.composite_variables
           summary['content'].append({
               'type': 'paragraph',
               'text': f"Analysis was performed using {len(variables_used)} variables: {', '.join(variables_used)}."
           })
       
       # Add top vulnerable wards
       if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
           top_wards = self.data_handler.vulnerability_rankings.sort_values('overall_rank').head(5)['WardName'].tolist()
           summary['content'].append({
               'type': 'paragraph',
               'text': f"The top 5 most vulnerable wards identified are: {', '.join(top_wards)}."
           })
           
           # Add vulnerability category counts
           if 'vulnerability_category' in self.data_handler.vulnerability_rankings.columns:
               category_counts = self.data_handler.vulnerability_rankings['vulnerability_category'].value_counts().to_dict()
               summary['content'].append({
                   'type': 'paragraph',
                   'text': f"Vulnerability categories: " + 
                         ', '.join([f"{cat}: {count} wards" for cat, count in category_counts.items()])
               })
       
       # Add urban extent summary
       if hasattr(self.data_handler, 'urban_extent_results') and self.data_handler.urban_extent_results:
           threshold_30 = self.data_handler.urban_extent_results.get(30)
           if threshold_30:
               total_wards = threshold_30['meets_threshold'] + threshold_30['below_threshold']
               urban_percentage = (threshold_30['meets_threshold']/total_wards*100) if total_wards > 0 else 0
               summary['content'].append({
                   'type': 'paragraph',
                   'text': f"At the 30% urban threshold, {threshold_30['meets_threshold']} wards ({urban_percentage:.1f}%) are classified as urban."
               })
               
               # Add information about non-urban high vulnerability wards
               non_urban_high_vuln = []
               if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
                   # Get top 10 vulnerable wards
                   top_vulnerable = self.data_handler.vulnerability_rankings.sort_values('overall_rank').head(10)
                   
                   # Check if these wards are urban at default threshold
                   results = self.data_handler.urban_extent_results[30]
                   for ward in top_vulnerable['WardName']:
                       if ward in results['below_threshold_wards']:
                           non_urban_high_vuln.append(ward)
               
               if non_urban_high_vuln:
                   summary['content'].append({
                       'type': 'paragraph',
                       'text': f"Notable finding: {len(non_urban_high_vuln)} high vulnerability wards are classified as non-urban, which may present logistical challenges for intervention planning."
                   })
       
       return summary
   
   def _generate_data_overview(self):
       """Generate data overview section"""
       overview = {
           'title': 'Data Overview',
           'content': []
       }
       
       # Add information about data sources
       overview['content'].append({
           'type': 'subheading',
           'text': 'Data Sources'
       })
       
       # Add CSV information
       if self.data_handler.csv_data is not None:
           csv_info = {
               'type': 'paragraph',
               'text': f"Data table: {len(self.data_handler.csv_data)} rows with {len(self.data_handler.csv_data.columns)} variables."
           }
           overview['content'].append(csv_info)
           
           # Add missing value information if available
           missing_columns = [col for col in self.data_handler.csv_data.columns 
                            if self.data_handler.csv_data[col].isna().sum() > 0]
           
           if missing_columns:
               missing_info = {
                   'type': 'paragraph',
                   'text': f"Data quality issues: {len(missing_columns)} variables had missing values that required imputation."
               }
               overview['content'].append(missing_info)
       
       # Add shapefile information
       if self.data_handler.shapefile_data is not None:
           shp_info = {
               'type': 'paragraph',
               'text': f"Geographical data: {len(self.data_handler.shapefile_data)} features with coordinate reference system {self.data_handler.shapefile_data.crs}."
           }
           overview['content'].append(shp_info)
       
       # Add variable statistics
       if self.data_handler.csv_data is not None:
           overview['content'].append({
               'type': 'subheading',
               'text': 'Variable Summary'
           })
           
           # Get variables used in analysis
           analysis_vars = self.data_handler.composite_variables if hasattr(self.data_handler, 'composite_variables') else []
           
           if analysis_vars:
               # Create variable statistics table
               var_stats = []
               for var in analysis_vars:
                   if var in self.data_handler.csv_data.columns:
                       values = self.data_handler.csv_data[var].dropna()
                       if len(values) > 0:
                           stat = {
                               'Variable': var,
                               'Full Name': get_full_variable_name(var),
                               'Min': f"{values.min():.2f}" if pd.api.types.is_numeric_dtype(values) else "N/A",
                               'Max': f"{values.max():.2f}" if pd.api.types.is_numeric_dtype(values) else "N/A",
                               'Mean': f"{values.mean():.2f}" if pd.api.types.is_numeric_dtype(values) else "N/A",
                               'Missing': f"{self.data_handler.csv_data[var].isna().sum()} ({self.data_handler.csv_data[var].isna().sum()/len(self.data_handler.csv_data)*100:.1f}%)"
                           }
                           var_stats.append(stat)
               
               if var_stats:
                   # Add variable statistics table
                   overview['content'].append({
                       'type': 'table',
                       'headers': ['Variable', 'Full Name', 'Min', 'Max', 'Mean', 'Missing'],
                       'rows': [[stat['Variable'], stat['Full Name'], stat['Min'], stat['Max'], stat['Mean'], stat['Missing']] for stat in var_stats]
                   })
       
       # Add data processing summary
       overview['content'].append({
           'type': 'subheading',
           'text': 'Data Processing Summary'
       })
       
       # Track analysis steps
       if self.metadata and hasattr(self.metadata, 'steps'):
           # Find cleaning and normalization steps
           cleaning_steps = [step for step in self.metadata.steps if 'clean' in step['step_name'].lower()]
           norm_steps = [step for step in self.metadata.steps if 'normal' in step['step_name'].lower()]
           
           if cleaning_steps:
               clean_info = {
                   'type': 'paragraph',
                   'text': f"Data cleaning included handling missing values in {len(missing_columns) if 'missing_columns' in locals() else '?'} variables."
               }
               overview['content'].append(clean_info)
           
           if norm_steps:
               norm_info = {
                   'type': 'paragraph',
                   'text': f"Data normalization was applied to create standardized scaled variables for analysis."
               }
               overview['content'].append(norm_info)
       
       return overview
   
   def _generate_methodology_section(self, detail_level):
       """
       Generate methodology section with appropriate detail level
       
       Args:
           detail_level: Level of detail (basic, standard, technical)
           
       Returns:
           dict: Methodology section data
       """
       methodology = {
           'title': 'Methodology',
           'content': []
       }
       
       # Add different content based on detail level
       if detail_level == 'basic':
           # Basic description of the methodology
           methodology['content'].append({
               'type': 'paragraph',
               'text': "The analysis uses a multi-step process to identify areas at highest risk for malaria transmission. " +
                       "Key steps include data cleaning, variable normalization, composite score calculation, and vulnerability ranking."
           })
           
           # Add information about variables
           if hasattr(self.data_handler, 'composite_variables') and self.data_handler.composite_variables:
               variables_used = self.data_handler.composite_variables
               var_text = "The analysis considered the following variables:\n\n"
               var_points = []
               for var in variables_used:
                   full_name = get_full_variable_name(var)
                   # Get relationship if available
                   relationship = "unknown"
                   if hasattr(self.data_handler, 'variable_relationships') and var in self.data_handler.variable_relationships:
                       relationship = self.data_handler.variable_relationships[var]
                   
                   rel_text = "higher = higher risk" if relationship == "direct" else "higher = lower risk" if relationship == "inverse" else "complex relationship with risk"
                   var_points.append(f"• {full_name} ({var}): {rel_text}")
               
               methodology['content'].append({
                   'type': 'paragraph',
                   'text': var_text + "\n".join(var_points)
               })
           
       elif detail_level == 'standard':
           # More detailed explanation of methods
           methodology['content'].append({
               'type': 'subheading',
               'text': 'Analysis Approach'
           })
           
           methodology['content'].append({
               'type': 'paragraph',
               'text': "The malaria risk analysis follows a standardized methodology with the following key steps:"
           })
           
           methodology['content'].append({
               'type': 'numbered_list',
               'items': [
                   "Data Cleaning: Missing values are imputed using spatial methods, mean values, or mode values as appropriate.",
                   "Variable Relationship Determination: Each variable's relationship with malaria risk (direct or inverse) is established.",
                   "Normalization: All variables are converted to a standard 0-1 scale based on their relationship with risk.",
                   "Composite Score Calculation: Multiple risk models are generated by combining different subsets of variables.",
                   "Vulnerability Ranking: Wards are ranked based on median risk scores across all models."
               ]
           })
           
           # Add variable relationship information
           if hasattr(self.data_handler, 'variable_relationships') and self.data_handler.variable_relationships:
               methodology['content'].append({
                   'type': 'subheading',
                   'text': 'Variable Relationships'
               })
               
               direct_vars = [var for var, rel in self.data_handler.variable_relationships.items() if rel == "direct"]
               inverse_vars = [var for var, rel in self.data_handler.variable_relationships.items() if rel == "inverse"]
               
               if direct_vars:
                   methodology['content'].append({
                       'type': 'paragraph',
                       'text': "Variables with direct relationship to malaria risk (higher values = higher risk):"
                   })
                   
                   methodology['content'].append({
                       'type': 'bullet_list',
                       'items': [f"{var} ({get_full_variable_name(var)})" for var in direct_vars]
                   })
               
               if inverse_vars:
                   methodology['content'].append({
                       'type': 'paragraph',
                       'text': "Variables with inverse relationship to malaria risk (higher values = lower risk):"
                   })
                   
                   methodology['content'].append({
                       'type': 'bullet_list',
                       'items': [f"{var} ({get_full_variable_name(var)})" for var in inverse_vars]
                   })
           
           # Add composite score model information
           if hasattr(self.data_handler, 'composite_scores') and self.data_handler.composite_scores:
               methodology['content'].append({
                   'type': 'subheading',
                   'text': 'Composite Score Models'
               })
               
               model_count = len(self.data_handler.composite_scores['model_formulas'])
               methodology['content'].append({
                   'type': 'paragraph',
                   'text': f"A total of {model_count} composite score models were generated using different combinations of variables."
               })
               
               # Example models
               if model_count > 0:
                   methodology['content'].append({
                       'type': 'paragraph',
                       'text': "Example model compositions:"
                   })
                   
                   model_examples = []
                   for i, model in enumerate(self.data_handler.composite_scores['model_formulas'][:3]):
                       model_examples.append(f"Model {i+1}: {', '.join(model['variables'])}")
                   
                   methodology['content'].append({
                       'type': 'bullet_list',
                       'items': model_examples
                   })
           
       else:  # technical detail level
           # Technical details of the methodology
           methodology['content'].append({
               'type': 'subheading',
               'text': 'Technical Methodology'
           })
           
           # Data cleaning technical details
           methodology['content'].append({
               'type': 'subheading',
               'text': 'Data Cleaning'
           })
           
           methodology['content'].append({
               'type': 'paragraph',
               'text': "Missing value imputation employed a hierarchical approach with the following methods (in order of priority):"
           })
           
           methodology['content'].append({
               'type': 'bullet_list',
               'items': [
                   "Spatial Neighbor Mean: Imputed values based on the average of adjacent wards. Used when possible to preserve spatial patterns.",
                   "Mean Imputation: For numeric variables without spatial context, values were replaced with the column mean.",
                   "Mode Imputation: For categorical variables, missing values were replaced with the most frequent value.",
                   "Forward/Backward Fill: As a last resort, values were propagated forward or backward."
               ]
           })
           
           # Add cleaning statistics if available
           if self.metadata and hasattr(self.metadata, 'calculations'):
               cleaning_calcs = [calc for calc in self.metadata.calculations 
                               if calc['operation'].startswith('imputation_')]
               
               if cleaning_calcs:
                   # Summarize imputation methods used
                   method_counts = {}
                   for calc in cleaning_calcs:
                       method = calc['operation'].replace('imputation_', '')
                       method_counts[method] = method_counts.get(method, 0) + 1
                   
                   methodology['content'].append({
                       'type': 'paragraph',
                       'text': "Imputation method distribution: " + 
                             ', '.join([f"{method}: {count} variables" for method, count in method_counts.items()])
                   })
           
           # Normalization technical details
           methodology['content'].append({
               'type': 'subheading',
               'text': 'Variable Normalization'
           })
           
           methodology['content'].append({
               'type': 'paragraph',
               'text': "Variables were normalized to a 0-1 scale using the following approaches:"
           })
           
           methodology['content'].append({
               'type': 'bullet_list',
               'items': [
                   "Direct relationship variables: Normalized using min-max scaling where normalized_value = (value - min) / (max - min)",
                   "Inverse relationship variables: First inverted using 1/value transformation, then min-max scaled",
                   "Special cases: Variables with all identical values were assigned a default value of 0.5"
               ]
           })
           
           # Composite score calculation technical details
           methodology['content'].append({
               'type': 'subheading',
               'text': 'Composite Score Calculation'
           })
           
           methodology['content'].append({
               'type': 'paragraph',
               'text': "Composite scores were calculated using the following approach:"
           })
           
           methodology['content'].append({
               'type': 'numbered_list',
               'items': [
                   "Generated all possible combinations of variables, from pairs to the full set",
                   "For each combination, calculated the arithmetic mean of normalized variable values",
                   "Scores range from 0 (lowest risk) to 1 (highest risk)",
                   "Multiple models ensure robustness against variable selection bias"
               ]
           })
           
           # Vulnerability ranking technical details
           methodology['content'].append({
               'type': 'subheading',
               'text': 'Vulnerability Ranking Methodology'
           })
           
           methodology['content'].append({
               'type': 'paragraph',
               'text': "Wards were ranked using the following process:"
           })
           
           methodology['content'].append({
               'type': 'numbered_list',
               'items': [
                   "Calculated median composite score across all models for each ward",
                   "Ranked wards by median score (highest score = highest vulnerability = rank 1)",
                   f"Categorized into {3 if hasattr(self.data_handler, 'vulnerability_rankings') else '?'} vulnerability categories based on rank terciles",
                   "High vulnerability: top third of ranks, Medium: middle third, Low: bottom third"
               ]
           })
           
           # Add technical information about urban analysis if available
           if hasattr(self.data_handler, 'urban_extent_results') and self.data_handler.urban_extent_results:
               methodology['content'].append({
                   'type': 'subheading',
                   'text': 'Urban Extent Analysis'
               })
               
               methodology['content'].append({
                   'type': 'paragraph',
                   'text': "Urban extent analysis was performed using the following approach:"
               })
               
               methodology['content'].append({
                   'type': 'numbered_list',
                   'items': [
                       "Applied multiple urban percentage thresholds (30%, 50%, 75%, 100%)",
                       "Classified wards as urban if their urban percentage equals or exceeds the threshold",
                       "Identified 'not ideal' situations where high vulnerability wards are classified as non-urban",
                       "These cases may present challenges for urban-focused intervention strategies"
                   ]
               })
       
       return methodology
   
   def _generate_results_section(self):
       """Generate results section with maps and visualizations"""
       results = {
           'title': 'Analysis Results',
           'content': []
       }
       
       # Add introduction to the results
       results['content'].append({
           'type': 'paragraph',
           'text': "The analysis produced several key outputs including composite risk maps, vulnerability rankings, and urban extent classifications. This section presents these results with visualizations and interpretations."
       })
       
       # Composite map visualization
       if hasattr(self.data_handler, 'composite_scores') and self.data_handler.composite_scores:
           results['content'].append({
               'type': 'subheading',
               'text': 'Composite Risk Maps'
           })
           
           results['content'].append({
               'type': 'paragraph',
               'text': "The composite risk maps show the spatial distribution of risk scores calculated from different combinations of variables. Darker colors indicate higher risk."
           })
           
           # Add composite map image if available
           composite_map_path = self._generate_composite_map_image()
           if composite_map_path:
               results['content'].append({
                   'type': 'image',
                   'path': composite_map_path,
                   'caption': 'Composite risk scores across different variable combinations. Darker colors indicate higher risk.'
               })
       
       # Vulnerability map visualization
       if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
           results['content'].append({
               'type': 'subheading',
               'text': 'Vulnerability Ranking'
           })
           
           results['content'].append({
               'type': 'paragraph',
               'text': "The vulnerability ranking map shows the classification of wards into vulnerability categories based on their median risk scores across all models."
           })
           
           # Add vulnerability map image if available
           vuln_map_path = self._generate_vulnerability_map_image()
           if vuln_map_path:
               results['content'].append({
                   'type': 'image',
                   'path': vuln_map_path,
                   'caption': 'Vulnerability ranking map. Wards are categorized as High (red), Medium (orange), or Low (yellow) vulnerability.'
               })
       
       # Box and whisker plot for distribution of ward scores
       results['content'].append({
           'type': 'subheading',
           'text': 'Risk Score Distribution'
       })
       
       results['content'].append({
           'type': 'paragraph',
           'text': "The box and whisker plot below shows the distribution of risk scores across different models for each ward. This visualization helps assess the consistency of risk scores and identify wards with high variability."
       })
       
       # Add box plot image if available
       box_plot_path = self._generate_box_plot_image()
       if box_plot_path:
           results['content'].append({
               'type': 'image',
               'path': box_plot_path,
               'caption': 'Box and whisker plot showing risk score distribution across models for each ward. Wards are ordered by overall vulnerability rank.'
           })
       
       return results
   
   def _generate_vulnerability_rankings(self):
       """Generate vulnerability rankings section"""
       rankings = {
           'title': 'Vulnerability Rankings',
           'content': []
       }
       
       # Add introduction to the rankings
       rankings['content'].append({
           'type': 'paragraph',
           'text': "This section presents the vulnerability rankings of all wards based on their median risk scores across composite models. Wards are categorized into High, Medium, and Low vulnerability groups to facilitate intervention prioritization."
       })
       
       # Add vulnerability statistics
       if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
           # Category counts
           if 'vulnerability_category' in self.data_handler.vulnerability_rankings.columns:
               category_counts = self.data_handler.vulnerability_rankings['vulnerability_category'].value_counts().to_dict()
               
               rankings['content'].append({
                   'type': 'paragraph',
                   'text': f"Vulnerability distribution: " + 
                         ', '.join([f"{cat}: {count} wards" for cat, count in category_counts.items()])
               })
           
           # Add top 10 most vulnerable wards
           rankings['content'].append({
               'type': 'subheading',
               'text': 'Most Vulnerable Wards'
           })
           
           top_wards = self.data_handler.vulnerability_rankings.sort_values('overall_rank').head(10)
           
           # Create table of top vulnerable wards
           top_wards_table = {
               'type': 'table',
               'headers': ['Rank', 'Ward Name', 'Risk Score', 'Category'],
               'rows': []
           }
           
           for _, row in top_wards.iterrows():
               # Get median score (which could be in different column names)
               if 'median_score' in row:
                   score = row['median_score']
               elif 'value' in row:
                   score = row['value']
               else:
                   score = 0.0
                   
               top_wards_table['rows'].append([
                   str(int(row['overall_rank'])),
                   row['WardName'],
                   f"{score:.3f}",
                   str(row['vulnerability_category']) if 'vulnerability_category' in row else 'Unknown'
               ])
           
           rankings['content'].append(top_wards_table)
           
           # Add bottom 10 least vulnerable wards
           rankings['content'].append({
               'type': 'subheading',
               'text': 'Least Vulnerable Wards'
           })
           
           bottom_wards = self.data_handler.vulnerability_rankings.sort_values('overall_rank', ascending=False).head(10)
           
           # Create table of least vulnerable wards
           bottom_wards_table = {
               'type': 'table',
               'headers': ['Rank', 'Ward Name', 'Risk Score', 'Category'],
               'rows': []
           }
           
           for _, row in bottom_wards.iterrows():
               # Get median score (which could be in different column names)
               if 'median_score' in row:
                   score = row['median_score']
               elif 'value' in row:
                   score = row['value']
               else:
                   score = 0.0
                   
               bottom_wards_table['rows'].append([
                   str(int(row['overall_rank'])),
                   row['WardName'],
                   f"{score:.3f}",
                   str(row['vulnerability_category']) if 'vulnerability_category' in row else 'Unknown'
               ])
           
           rankings['content'].append(bottom_wards_table)
       
       return rankings
   
   def _generate_urban_extent_section(self):
       """Generate urban extent analysis section"""
       urban = {
           'title': 'Urban Extent Analysis',
           'content': []
       }
       
       # Add introduction to urban extent analysis
       urban['content'].append({
           'type': 'paragraph',
           'text': "Urban extent analysis examines the relationship between urban areas and malaria vulnerability. This analysis is crucial for intervention planning as urban and rural areas may require different approaches."
       })
       
       # Add urban threshold results
       if hasattr(self.data_handler, 'urban_extent_results') and self.data_handler.urban_extent_results:
           urban['content'].append({
               'type': 'subheading',
               'text': 'Urban Classification by Threshold'
           })
           
           # Create urban threshold table
           threshold_table = {
               'type': 'table',
               'headers': ['Threshold', 'Urban Wards', 'Non-Urban Wards', 'Urban Percentage'],
               'rows': []
           }
           
           for threshold, results in sorted(self.data_handler.urban_extent_results.items()):
               total_wards = results['meets_threshold'] + results['below_threshold']
               urban_percentage = (results['meets_threshold'] / total_wards * 100) if total_wards > 0 else 0
               
               threshold_table['rows'].append([
                   f"{threshold}%",
                   str(results['meets_threshold']),
                   str(results['below_threshold']),
                   f"{urban_percentage:.1f}%"
               ])
           
           urban['content'].append(threshold_table)
           
           # Add urban extent map
           urban['content'].append({
               'type': 'subheading',
               'text': 'Urban Extent Map'
           })
           
           urban['content'].append({
               'type': 'paragraph',
               'text': "The map below shows the urban extent at the 30% threshold overlaid with vulnerability rankings. Urban areas are colored by vulnerability level, while non-urban areas are shown in gray."
           })
           
           # Add urban extent map image if available
           urban_map_path = self._generate_urban_extent_map_image()
           if urban_map_path:
               urban['content'].append({
                   'type': 'image',
                   'path': urban_map_path,
                   'caption': 'Urban extent map at 30% threshold. Urban areas are colored by vulnerability level, non-urban areas are gray.'
               })
           
           # Add information about non-urban high vulnerability wards
           non_urban_high_vuln = []
           if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
               # Get top 10 vulnerable wards
               top_vulnerable = self.data_handler.vulnerability_rankings.sort_values('overall_rank').head(10)
               
               # Default threshold for urban/rural classification
               default_threshold = 30
               
               # Check if these wards are urban at default threshold
               if default_threshold in self.data_handler.urban_extent_results:
                   results = self.data_handler.urban_extent_results[default_threshold]
                   for ward in top_vulnerable['WardName']:
                       if ward in results['below_threshold_wards']:
                           non_urban_high_vuln.append(ward)
           
           if non_urban_high_vuln:
               urban['content'].append({
                   'type': 'subheading',
                   'text': 'Non-Urban High Vulnerability Wards'
               })
               
               urban['content'].append({
                   'type': 'paragraph',
                   'text': "The following high vulnerability wards are classified as non-urban (below 30% threshold). These wards may present logistical challenges for urban-focused interventions and should be considered carefully in planning."
               })
               
               urban['content'].append({
                   'type': 'bullet_list',
                   'items': non_urban_high_vuln
               })
       
       return urban
   
   def _generate_recommendations(self):
       """Generate recommendations section"""
       recommendations = {
           'title': 'Recommendations',
           'content': []
       }
       
       # Add introduction to recommendations
       recommendations['content'].append({
           'type': 'paragraph',
           'text': "Based on the analysis results, the following recommendations are provided for malaria intervention planning. These recommendations aim to optimize resource allocation and intervention strategy selection."
       })
       
       # Geographical prioritization recommendations
       recommendations['content'].append({
           'type': 'subheading',
           'text': 'Geographical Prioritization'
       })
       
       # Get top vulnerable wards for recommendations
       top_wards = []
       if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
           top_wards = self.data_handler.vulnerability_rankings.sort_values('overall_rank').head(5)['WardName'].tolist()
       
       # Add prioritization recommendations
       priority_recs = [
           f"Focus initial interventions on highest vulnerability wards: {', '.join(top_wards)}." if top_wards else 
           "Focus initial interventions on highest vulnerability wards as identified in the rankings.",
           "Consider clustering of high vulnerability areas for efficient resource deployment.",
           "Balance urban and rural interventions based on the urban extent analysis results."
       ]
       
       recommendations['content'].append({
           'type': 'bullet_list',
           'items': priority_recs
       })
       
       # Intervention strategy recommendations
       recommendations['content'].append({
           'type': 'subheading',
           'text': 'Intervention Strategy'
       })
       
       # Get variables information for strategy recommendations
       variables_used = self.data_handler.composite_variables if hasattr(self.data_handler, 'composite_variables') else []
       
       strategy_recs = []
       
       # Add strategy recommendations based on variables
       if any(var in ['rainfall', 'precipitation', 'mean_rainfall'] for var in variables_used):
           strategy_recs.append("Increase vector control measures during and after rainy seasons in high-risk areas.")
       
       if any(var in ['ndvi', 'evi', 'vegetation', 'mean_ndvi', 'mean_evi'] for var in variables_used):
           strategy_recs.append("Target areas with high vegetation indices for larval source management.")
       
       if any(var in ['distance_to_water', 'water_dist'] for var in variables_used):
           strategy_recs.append("Implement enhanced surveillance near water bodies in high vulnerability wards.")
       
       if any(var in ['housing_quality', 'house_qual'] for var in variables_used):
           strategy_recs.append("Consider housing improvement programs in high vulnerability areas with poor housing quality.")
       
       if any(var in ['population', 'pop_density', 'population_density'] for var in variables_used):
           strategy_recs.append("Scale intervention coverage based on population density to maximize impact.")
       
       # Add general recommendations if specific ones couldn't be generated
       if not strategy_recs:
           strategy_recs = [
               "Tailor intervention strategies based on the specific risk factors in each area.",
               "Consider seasonal timing of interventions based on climate factors.",
               "Implement comprehensive intervention packages in highest risk areas.",
               "Develop targeted strategies for non-urban high vulnerability wards."
           ]
       
       recommendations['content'].append({
           'type': 'bullet_list',
           'items': strategy_recs
       })
       
       # Monitoring and evaluation recommendations
       recommendations['content'].append({
           'type': 'subheading',
           'text': 'Monitoring and Evaluation'
       })
       
       me_recs = [
           "Establish baseline surveys in priority wards before intervention implementation.",
           "Implement regular monitoring of intervention coverage and impact.",
           "Consider longitudinal studies in selected high and low vulnerability areas to validate risk models.",
           "Update analysis periodically as new data becomes available to track changes in vulnerability patterns."
       ]
       
       recommendations['content'].append({
           'type': 'bullet_list',
           'items': me_recs
       })
       
       return recommendations
   
   def _generate_appendix(self, detail_level):
       """
       Generate appendix with technical details
       
       Args:
           detail_level: Level of detail (basic, standard, technical)
           
       Returns:
           dict: Appendix section
       """
       appendix = {
           'title': 'Appendix',
           'content': []
       }
       
       # Only include full technical details for technical detail level
       if detail_level == 'technical':
           # Add analysis steps details
           if self.metadata and hasattr(self.metadata, 'steps'):
               appendix['content'].append({
                   'type': 'subheading',
                   'text': 'Detailed Analysis Steps'
               })
               
               # Create table of analysis steps
               steps_table = {
                   'type': 'table',
                   'headers': ['Step', 'Algorithm', 'Execution Time (s)', 'Parameters'],
                   'rows': []
               }
               
               for step in self.metadata.steps:
                   if 'step_name' in step and 'algorithm' in step:
                       # Format parameters nicely
                       params = step.get('parameters', {})
                       if params:
                           params_str = ', '.join([f"{k}: {v}" for k, v in params.items()])
                       else:
                           params_str = "None"
                       
                       steps_table['rows'].append([
                           step['step_name'],
                           step['algorithm'],
                           f"{step.get('execution_time', 0):.2f}",
                           params_str
                       ])
               
               appendix['content'].append(steps_table)
           
           # Add variable relationship details
           if hasattr(self.data_handler, 'variable_relationships') and self.data_handler.variable_relationships:
               appendix['content'].append({
                   'type': 'subheading',
                   'text': 'Variable Relationship Details'
               })
               
               # Create table of variable relationships
               relationship_table = {
                   'type': 'table',
                   'headers': ['Variable', 'Full Name', 'Relationship', 'Normalization'],
                   'rows': []
               }
               
               for var, relationship in self.data_handler.variable_relationships.items():
                   full_name = get_full_variable_name(var)
                   norm_method = "Min-Max Scaling" if relationship == "direct" else "Inversion + Min-Max" if relationship == "inverse" else "Unknown"
                   
                   relationship_table['rows'].append([
                       var,
                       full_name,
                       relationship.capitalize(),
                       norm_method
                   ])
               
               appendix['content'].append(relationship_table)
           
           # Add details about composite score models
           if hasattr(self.data_handler, 'composite_scores') and self.data_handler.composite_scores:
               appendix['content'].append({
                   'type': 'subheading',
                   'text': 'Composite Score Model Details'
               })
               
               # Create table of models
               model_table = {
                   'type': 'table',
                   'headers': ['Model', 'Variables Used'],
                   'rows': []
               }
               
               for model in self.data_handler.composite_scores['model_formulas']:
                   model_table['rows'].append([
                       model['model'],
                       ', '.join(model['variables'])
                   ])
               
               appendix['content'].append(model_table)
       
       # Include complete ward rankings for both standard and technical
       if detail_level in ['standard', 'technical']:
           if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
               appendix['content'].append({
                   'type': 'subheading',
                   'text': 'Complete Ward Rankings'
               })
               
               # Create table of all ward rankings
               rankings_table = {
                   'type': 'table',
                   'headers': ['Rank', 'Ward Name', 'Risk Score', 'Category'],
                   'rows': []
               }
               
               # Sort wards by rank
               sorted_rankings = self.data_handler.vulnerability_rankings.sort_values('overall_rank')
               
               for _, row in sorted_rankings.iterrows():
                   # Get median score (which could be in different column names)
                   if 'median_score' in row:
                       score = row['median_score']
                   elif 'value' in row:
                       score = row['value']
                   else:
                       score = 0.0
                       
                   rankings_table['rows'].append([
                       str(int(row['overall_rank'])),
                       row['WardName'],
                       f"{score:.3f}",
                       str(row['vulnerability_category']) if 'vulnerability_category' in row else 'Unknown'
                   ])
               
               appendix['content'].append(rankings_table)
       
       return appendix
   
   def _generate_custom_section(self, section_name):
       """
       Generate a custom section based on name
       
       Args:
           section_name: Name of the custom section
           
       Returns:
           dict: Custom section data
       """
       section = {
           'title': section_name.replace('_', ' ').title(),
           'content': []
       }
       
       # Get LLM to generate content for this section
       if self.llm_manager:
           llm_context = {
               'metadata': self._get_report_metadata(),
               'section_name': section_name,
               'data_handler': {
                   'has_vulnerability_rankings': hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None,
                   'has_urban_extent': hasattr(self.data_handler, 'urban_extent_results') and self.data_handler.urban_extent_results,
                   'variables_used': self.data_handler.composite_variables if hasattr(self.data_handler, 'composite_variables') else []
               }
           }
           
           prompt = f"Generate content for a custom report section titled '{section_name.replace('_', ' ').title()}'. Include relevant analysis insights and recommendations."
           
           content = self.llm_manager.generate_response(
               prompt=prompt,
               context=llm_context,
               session_id=self.session_id
           )
           
           if content:
               # Add the content as a paragraph
               section['content'].append({
                   'type': 'paragraph',
                   'text': content
               })
           else:
               section['content'].append({
                   'type': 'paragraph',
                   'text': f"No content available for {section_name.replace('_', ' ').title()} section."
               })
       else:
           section['content'].append({
               'type': 'paragraph',
               'text': f"Custom section: {section_name.replace('_', ' ').title()}.\nLLM manager not available to generate content."
           })
       
       return section
   
   def _generate_composite_map_image(self):
       """Generate and save a composite map image"""
       try:
           # Check if we have the necessary data
           if not hasattr(self.data_handler, 'composite_scores') or not self.data_handler.composite_scores:
               logger.warning("No composite scores available for map generation")
               return None
           
           if not hasattr(self.data_handler, 'shapefile_data') or self.data_handler.shapefile_data is None:
               logger.warning("No shapefile data available for map generation")
               return None
           
           # Create figure with plotly
           import plotly.graph_objects as go
           from plotly.subplots import make_subplots
           
           # Determine how many models to show (max 4)
           model_columns = [col for col in self.data_handler.composite_scores['scores'].columns if col.startswith('model_')]
           models_to_show = min(4, len(model_columns))
           
           if models_to_show == 0:
               logger.warning("No model columns found in composite scores")
               return None
           
           # Determine grid layout
           if models_to_show == 1:
               rows, cols = 1, 1
           elif models_to_show == 2:
               rows, cols = 1, 2
           else:
               rows, cols = 2, 2
           
           # Create subplot titles
           subplot_titles = []
           for i, model in enumerate(model_columns[:models_to_show]):
               model_idx = i + 1
               variables = self.data_handler.composite_scores['model_formulas'][i]['variables'] if i < len(self.data_handler.composite_scores['model_formulas']) else []
               if variables:
                   # Create title with variables
                   var_display = ', '.join(variables[:3])
                   if len(variables) > 3:
                       var_display += f" +{len(variables)-3} more"
                   title = f"Model {model_idx}: {var_display}"
               else:
                   title = f"Model {model_idx}"
               subplot_titles.append(title)
           
           # Create subplots
           fig = make_subplots(
               rows=rows,
               cols=cols,
               specs=[[{"type": "choropleth"}] * cols for _ in range(rows)],
               subplot_titles=subplot_titles
           )
           
           # Get shapefile data
           shapefile_data = self.data_handler.shapefile_data.copy()
           
           # Combine with composite scores
           combined_data = shapefile_data.merge(
               self.data_handler.composite_scores['scores'],
               on='WardName',
               how='left'
           )
           
           # Convert to geojson for plotting
           import json
           geojson = json.loads(combined_data.to_json())
           
           # Add each model as a subplot
           for i, model in enumerate(model_columns[:models_to_show]):
               row = i // cols + 1
               col = i % cols + 1
               
               # Create choropleth
               fig.add_trace(
                   go.Choropleth(
                       geojson=geojson,
                       locations=combined_data.index,
                       z=combined_data[model],
                       colorscale='YlOrRd',
                       zmin=0,
                       zmax=1,
                       marker_line_width=0.5,
                       marker_line_color='black',
                       colorbar=dict(
                           title="Risk Score",
                           tickvals=[0, 0.5, 1],
                           ticktext=["Low", "Medium", "High"]
                       ) if i == 0 else None,
                       showscale=i == 0,
                       hovertemplate='<b>%{text}</b><br>Risk Score: %{z:.3f}<extra></extra>',
                       text=combined_data['WardName']
                   ),
                   row=row, col=col
               )
           
           # Update layout
           fig.update_layout(
               title="Composite Risk Maps",
               height=600,
               width=800,
               margin=dict(t=60, b=40, l=40, r=40)
           )
           
           # Save the figure
           img_filename = f"composite_map_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
           img_path = os.path.join(self.report_folder, img_filename)
           fig.write_image(img_path)
           
           return img_path
       
       except Exception as e:
           logger.error(f"Error generating composite map image: {str(e)}")
           return None
   
   def _generate_vulnerability_map_image(self):
       """Generate and save a vulnerability map image"""
       try:
           # Check if we have the necessary data
           if not hasattr(self.data_handler, 'vulnerability_rankings') or self.data_handler.vulnerability_rankings is None:
               logger.warning("No vulnerability rankings available for map generation")
               return None
           
           if not hasattr(self.data_handler, 'shapefile_data') or self.data_handler.shapefile_data is None:
               logger.warning("No shapefile data available for map generation")
               return None
           
           # Create figure with plotly
           import plotly.graph_objects as go
           
           # Get shapefile data
           shapefile_data = self.data_handler.shapefile_data.copy()
           
           # Combine with vulnerability rankings
           combined_data = shapefile_data.merge(
               self.data_handler.vulnerability_rankings,
               on='WardName',
               how='left'
           )
           
           # Get color values based on vulnerability category
           color_values = []
           if 'vulnerability_category' in combined_data.columns:
               # Map categories to numeric values
               category_map = {'High': 3, 'Medium': 2, 'Low': 1, None: 0}
               color_values = [category_map.get(cat, 0) for cat in combined_data['vulnerability_category']]
               colorscale = [[0, 'rgba(220,220,220,0.5)'], [0.25, '#d7191c'], [0.5, '#fdae61'], [0.75, '#ffffbf']]
               z_values = color_values
               colorbar_title = "Vulnerability"
               tickvals = [1.5, 2.5, 3.5]
               ticktext = ['Low', 'Medium', 'High']
           else:
               # Use overall rank
               z_values = combined_data['overall_rank'] if 'overall_rank' in combined_data.columns else None
               colorscale = 'Plasma_r'  # Reverse plasma so high vulnerability (low rank) is dark
               colorbar_title = "Vulnerability Rank"
               max_rank = combined_data['overall_rank'].max() if 'overall_rank' in combined_data.columns else 100
               tickvals = [1, max_rank / 2, max_rank]
               ticktext = ['High', 'Medium', 'Low']
           
           # Convert to geojson for plotting
           import json
           geojson = json.loads(combined_data.to_json())
           
           # Create figure
           fig = go.Figure()
           
           # Add choropleth
           fig.add_trace(
               go.Choropleth(
                   geojson=geojson,
                   locations=combined_data.index,
                   z=z_values,
                   colorscale=colorscale,
                   marker_line_width=0.5,
                   marker_line_color='black',
                   colorbar=dict(
                       title=colorbar_title,
                       tickvals=tickvals,
                       ticktext=ticktext
                   ),
                   hovertemplate='<b>%{text}</b><br>Rank: %{customdata}<extra></extra>',
                   text=combined_data['WardName'],
                   customdata=combined_data['overall_rank'] if 'overall_rank' in combined_data.columns else None
               )
           )
           
           # Update layout
           fig.update_layout(
               title="Vulnerability Ranking Map",
               height=600,
               width=800,
               margin=dict(t=60, b=40, l=40, r=40)
           )
           
           # Save the figure
           img_filename = f"vulnerability_map_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
           img_path = os.path.join(self.report_folder, img_filename)
           fig.write_image(img_path)
           
           return img_path
       
       except Exception as e:
           logger.error(f"Error generating vulnerability map image: {str(e)}")
           return None
   
   def _generate_box_plot_image(self):
       """Generate and save a box plot image"""
       try:
           # Check if we have the necessary data
           if not hasattr(self.data_handler, 'composite_scores') or not self.data_handler.composite_scores:
               logger.warning("No composite scores available for box plot generation")
               return None
           
           # Get scores and model columns
           scores_df = self.data_handler.composite_scores['scores']
           model_cols = [col for col in scores_df.columns if col.startswith('model_')]
           
           if not model_cols:
               logger.warning("No model columns found in composite scores")
               return None
           
           # Get vulnerability rankings
           vulnerability_rankings = None
           if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
               vulnerability_rankings = self.data_handler.vulnerability_rankings
           
           # Create figure with plotly
           import plotly.graph_objects as go
           import pandas as pd
           
           # Melt the dataframe for box plot
           melted_df = pd.melt(
               scores_df,
               id_vars=['WardName'],
               value_vars=model_cols,
               var_name='Model',
               value_name='Risk Score'
           )
           
           # Join with vulnerability rankings if available
           if vulnerability_rankings is not None:
               melted_df = melted_df.merge(
                   vulnerability_rankings[['WardName', 'overall_rank', 'vulnerability_category']],
                   on='WardName',
                   how='left'
               )
               
               # Sort by overall rank
               melted_df = melted_df.sort_values('overall_rank')
           
           # Only show top 20 most vulnerable wards
           if vulnerability_rankings is not None:
               top_wards = vulnerability_rankings.sort_values('overall_rank').head(20)['WardName'].tolist()
               melted_df = melted_df[melted_df['WardName'].isin(top_wards)]
           else:
               # Limit to first 20 wards
               unique_wards = melted_df['WardName'].unique()
               if len(unique_wards) > 20:
                   top_wards = unique_wards[:20]
                   melted_df = melted_df[melted_df['WardName'].isin(top_wards)]
           
           # Create categorical y-axis with ward names
           ward_order = melted_df['WardName'].unique()
           
           # Create figure
           fig = go.Figure()
           
           # Add box plots for each ward
           for ward in ward_order:
               ward_data = melted_df[melted_df['WardName'] == ward]
               
               # Determine color based on vulnerability category if available
               color = '#69b3a2'  # Default color
               if 'vulnerability_category' in ward_data.columns:
                   category = ward_data['vulnerability_category'].iloc[0] if not ward_data.empty else None
                   if category == 'High':
                       color = '#d7191c'
                   elif category == 'Medium':
                       color = '#fdae61'
                   elif category == 'Low':
                       color = '#ffffbf'
               
               fig.add_trace(
                   go.Box(
                       x=ward_data['Risk Score'],
                       name=ward,
                       marker_color=color,
                       orientation='h',
                       boxmean=True
                   )
               )
           
           # Update layout
           fig.update_layout(
               title="Risk Score Distribution Across Models (Top 20 Vulnerable Wards)",
               xaxis_title="Risk Score",
               yaxis_title="Ward Name",
               yaxis={'categoryorder': 'array', 'categoryarray': list(reversed(ward_order))},
               height=800,
               width=800,
               margin=dict(t=60, b=60, l=200, r=40)
           )
           
           # Save the figure
           img_filename = f"box_plot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
           img_path = os.path.join(self.report_folder, img_filename)
           fig.write_image(img_path)
           
           return img_path
       
       except Exception as e:
           logger.error(f"Error generating box plot image: {str(e)}")
           return None
   
   def _generate_urban_extent_map_image(self):
       """Generate and save an urban extent map image"""
       try:
           # Check if we have the necessary data
           if not hasattr(self.data_handler, 'urban_extent_results') or not self.data_handler.urban_extent_results:
               logger.warning("No urban extent results available for map generation")
               return None
           
           if not hasattr(self.data_handler, 'shapefile_data') or self.data_handler.shapefile_data is None:
               logger.warning("No shapefile data available for map generation")
               return None
           
           # Create figure with plotly
           import plotly.graph_objects as go
           
           # Get shapefile data
           shapefile_data = self.data_handler.shapefile_data.copy()
           
           # Default threshold for urban/rural classification
           default_threshold = 30
           
           # Check if default threshold exists
           if default_threshold not in self.data_handler.urban_extent_results:
               thresholds = list(self.data_handler.urban_extent_results.keys())
               if thresholds:
                   default_threshold = thresholds[0]
               else:
                   logger.warning("No urban thresholds available")
                   return None
           
           # Get urban extent results
           urban_results = self.data_handler.urban_extent_results[default_threshold]
           
           # Combine with shapefile data
           # Create a column for urban classification
           urban_wards = urban_results['meets_threshold_wards']
           shapefile_data['is_urban'] = shapefile_data['WardName'].isin(urban_wards)
           
           # Add vulnerability data if available
           if hasattr(self.data_handler, 'vulnerability_rankings') and self.data_handler.vulnerability_rankings is not None:
               combined_data = shapefile_data.merge(
                   self.data_handler.vulnerability_rankings,
                   on='WardName',
                   how='left'
               )
           else:
               combined_data = shapefile_data
           
           # Convert to geojson for plotting
           import json
           geojson = json.loads(combined_data.to_json())
           
           # Create figure
           fig = go.Figure()
           
           # Create two traces: one for urban areas (colored by vulnerability) and one for non-urban (gray)
           
           # 1. Urban areas
           urban_data = combined_data[combined_data['is_urban']]
           
           if not urban_data.empty:
               # Determine coloring based on vulnerability if available
               if 'vulnerability_category' in urban_data.columns:
                   # Map categories to numeric values
                   category_map = {'High': 3, 'Medium': 2, 'Low': 1, None: 0}
                   color_values = [category_map.get(cat, 0) for cat in urban_data['vulnerability_category']]
                   colorscale = [[0, 'rgba(220,220,220,0.5)'], [0.25, '#d7191c'], [0.5, '#fdae61'], [0.75, '#ffffbf']]
                   z_values = color_values
                   colorbar_title = "Vulnerability"
                   tickvals = [1.5, 2.5, 3.5]
                   ticktext = ['Low', 'Medium', 'High']
               elif 'overall_rank' in urban_data.columns:
                   # Use overall rank
                   z_values = urban_data['overall_rank']
                   colorscale = 'Plasma_r'  # Reverse plasma so high vulnerability (low rank) is dark
                   colorbar_title = "Vulnerability Rank"
                   max_rank = urban_data['overall_rank'].max() if not urban_data.empty else 100
                   tickvals = [1, max_rank / 2, max_rank]
                   ticktext = ['High', 'Medium', 'Low']
               else:
                   # Default to urban percentage if available
                   urban_col = None
                   for col in ['UrbanPercentage', 'UrbanPercent', 'Urban_Percent', 'urbanPercent']:
                       if col in urban_data.columns:
                           urban_col = col
                           break
                   
                   if urban_col:
                       z_values = urban_data[urban_col]
                       colorscale = 'YlOrRd'
                       colorbar_title = "Urban Percentage"
                       tickvals = [default_threshold, 50, 100]
                       ticktext = [f"{default_threshold}%", "50%", "100%"]
                   else:
                       # Fallback to constant value
                       z_values = [1] * len(urban_data)
                       colorscale = [[0, '#d7191c'], [1, '#d7191c']]
                       colorbar_title = "Urban Areas"
                       tickvals = [0.5]
                       ticktext = ["Urban"]
               
               fig.add_trace(
                   go.Choropleth(
                       geojson=geojson,
                       locations=urban_data.index,
                       z=z_values,
                       colorscale=colorscale,
                       marker_line_width=0.5,
                       marker_line_color='black',
                       colorbar=dict(
                           title=colorbar_title,
                           tickvals=tickvals,
                           ticktext=ticktext
                       ),
                       hovertemplate='<b>%{text}</b><br>Urban Area<extra></extra>',
                       text=urban_data['WardName'],
                       name='Urban Areas'
                   )
               )
           
           # 2. Non-urban areas (gray)
           non_urban_data = combined_data[~combined_data['is_urban']]
           
           if not non_urban_data.empty:
               fig.add_trace(
                   go.Choropleth(
                       geojson=geojson,
                       locations=non_urban_data.index,
                       z=[0] * len(non_urban_data),  # Constant value for uniform color
                       colorscale=[[0, 'rgba(200,200,200,0.4)'], [1, 'rgba(200,200,200,0.4)']],  # Light gray
                       marker_line_width=0.5,
                       marker_line_color='rgba(150,150,150,0.3)',
                       showscale=False,
                       hovertemplate='<b>%{text}</b><br>Non-Urban Area<extra></extra>',
                       text=non_urban_data['WardName'],
                       name='Non-Urban Areas'
                   )
               )
           
           # 3. Add blue outlines for non-urban high vulnerability wards
           non_urban_high_vuln = []
           if 'vulnerability_category' in combined_data.columns and 'is_urban' in combined_data.columns:
               # Find non-urban wards with high vulnerability
               non_urban_high_data = combined_data[
                   (~combined_data['is_urban']) & 
                   (combined_data['vulnerability_category'] == 'High')
               ]
               
               if not non_urban_high_data.empty:
                   non_urban_high_vuln = non_urban_high_data['WardName'].tolist()
                   
                   fig.add_trace(
                       go.Choropleth(
                           geojson=geojson,
                           locations=non_urban_high_data.index,
                           z=[1] * len(non_urban_high_data),  # Constant value
                           colorscale=[[0, 'rgba(0,0,0,0)'], [1, 'rgba(0,0,0,0)']],  # Transparent fill
                           marker_line_width=3,
                           marker_line_color='blue',
                           showscale=False,
                           hovertemplate='<b>%{text}</b><br>Non-Urban High Vulnerability Ward<extra></extra>',
                           text=non_urban_high_data['WardName'],
                           name='Non-Urban High Vulnerability'
                       )
                   )
           
           # Update layout
           title = f"Urban Extent Map ({default_threshold}% Threshold)"
           if non_urban_high_vuln:
               title += f"<br><span style='font-size:14px'>Blue outlines: Non-urban high vulnerability wards</span>"
           
           fig.update_layout(
               title=title,
               height=600,
               width=800,
               margin=dict(t=60, b=40, l=40, r=40)
           )
           
           # Save the figure
           img_filename = f"urban_extent_map_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
           img_path = os.path.join(self.report_folder, img_filename)
           fig.write_image(img_path)
           
           return img_path
       
       except Exception as e:
           logger.error(f"Error generating urban extent map image: {str(e)}")
           return None
   
   def _generate_pdf_report(self, report_data):
       """
       Generate PDF report
       
       Args:
           report_data: Data for the report
           
       Returns:
           tuple: (report_path, report_url)
       """
       try:
           # Create unique filename
           timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
           filename = f"malaria_risk_report_{timestamp}.pdf"
           report_path = os.path.join(self.report_folder, filename)
           
           # Generate HTML content first
           html_path, _ = self._generate_html_report(report_data, save_file=False)
           
           # Convert HTML to PDF using weasyprint
           try:
               from weasyprint import HTML
               HTML(html_path).write_pdf(report_path)
           except ImportError:
               # Fallback to alternative PDF generation if weasyprint not available
               from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
               from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
               from reportlab.lib import colors
               from reportlab.lib.pagesizes import letter
               from reportlab.lib.units import inch
               
               # Create document
               doc = SimpleDocTemplate(report_path, pagesize=letter)
               styles = getSampleStyleSheet()
               
               # Create custom styles
               styles.add(ParagraphStyle(
                   name='Title',
                   parent=styles['Heading1'],
                   fontSize=24,
                   spaceAfter=12,
                   textColor=self.format_config['color_primary']
               ))
               
               styles.add(ParagraphStyle(
                   name='Subtitle',
                   parent=styles['Heading2'],
                   fontSize=16,
                   spaceAfter=12,
                   textColor=self.format_config['color_secondary']
               ))
               
               styles.add(ParagraphStyle(
                   name='SectionHeading',
                   parent=styles['Heading2'],
                   fontSize=14,
                   spaceAfter=10,
                   textColor=self.format_config['color_primary']
               ))
               
               styles.add(ParagraphStyle(
                   name='SubsectionHeading',
                   parent=styles['Heading3'],
                   fontSize=12,
                   spaceAfter=8,
                   textColor=self.format_config['color_secondary']
               ))
               
               # Build content
               content = []
               
               # Title
               content.append(Paragraph(report_data['title'], styles['Title']))
               content.append(Paragraph(report_data['subtitle'], styles['Subtitle']))
               content.append(Paragraph(f"Generated: {report_data['date']}", styles['Normal']))
               content.append(Spacer(1, 0.25*inch))
               
               # Add each section
               for section_name, section in report_data['sections'].items():
                   # Section heading
                   content.append(Paragraph(section['title'], styles['SectionHeading']))
                   
                   # Add section content
                   for item in section['content']:
                       if item['type'] == 'paragraph':
                           content.append(Paragraph(item['text'], styles['Normal']))
                           content.append(Spacer(1, 0.1*inch))
                       
                       elif item['type'] == 'subheading':
                           content.append(Paragraph(item['text'], styles['SubsectionHeading']))
                       
                       elif item['type'] == 'bullet_list':
                           for bullet in item['items']:
                               content.append(Paragraph(f"• {bullet}", styles['Normal']))
                           content.append(Spacer(1, 0.1*inch))
                       
                       elif item['type'] == 'numbered_list':
                           for i, entry in enumerate(item['items']):
                               content.append(Paragraph(f"{i+1}. {entry}", styles['Normal']))
                           content.append(Spacer(1, 0.1*inch))
                       
                       elif item['type'] == 'table':
                           # Create table data
                           table_data = [item['headers']]
                           for row in item['rows']:
                               table_data.append(row)
                           
                           # Create table
                           table = Table(table_data)
                           
                           # Add table style
                           table.setStyle(TableStyle([
                               ('BACKGROUND', (0, 0), (-1, 0), self.format_config['color_primary']),
                               ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                               ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                               ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                               ('FONTSIZE', (0, 0), (-1, 0), 10),
                               ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                               ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                               ('GRID', (0, 0), (-1, -1), 1, colors.black),
                           ]))
                           
                           content.append(table)
                           content.append(Spacer(1, 0.2*inch))
                       
                       elif item['type'] == 'image':
                           if os.path.exists(item['path']):
                               img = Image(item['path'], width=6*inch, height=4*inch)
                               content.append(img)
                               if 'caption' in item:
                                   content.append(Paragraph(item['caption'], styles['Caption']))
                               content.append(Spacer(1, 0.2*inch))
                   
                   # Add spacer between sections
                   content.append(Spacer(1, 0.5*inch))
               
               # Build the PDF
               doc.build(content)
           
           # Create report URL
           report_url = f"/download_report/{filename}"
           
           return report_path, report_url
           
       except Exception as e:
           logger.error(f"Error generating PDF report: {str(e)}", exc_info=True)
           raise
   
   def _generate_html_report(self, report_data, save_file=True):
       """
       Generate HTML report
       
       Args:
           report_data: Data for the report
           save_file: Whether to save the file (True) or just return content (False)
           
       Returns:
           tuple: (report_path, report_url)
       """
       try:
           # Start building HTML content
           html_content = f"""
           <!DOCTYPE html>
           <html lang="en">
           <head>
               <meta charset="UTF-8">
               <meta name="viewport" content="width=device-width, initial-scale=1.0">
               <title>{report_data['title']}</title>
               <style>
                   body {{
                       font-family: 'Arial', sans-serif;
                       line-height: 1.6;
                       color: {self.format_config['color_text']};
                       max-width: {self.format_config['max_width']}px;
                       margin: 0 auto;
                       padding: {self.format_config['page_margin']}px;
                   }}
                   h1 {{
                       color: {self.format_config['color_primary']};
                       font-size: 32px;
                       margin-bottom: 10px;
                   }}
                   h2 {{
                       color: {self.format_config['color_primary']};
                       font-size: 24px;
                       border-bottom: 1px solid #ddd;
                       padding-bottom: 10px;
                       margin-top: 30px;
                   }}
                   h3 {{
                       color: {self.format_config['color_secondary']};
                       font-size: 18px;
                   }}
                   table {{
                       border-collapse: collapse;
                       width: 100%;
                       margin: 20px 0;
                   }}
                   th, td {{
                       border: 1px solid #ddd;
                       padding: 8px;
                       text-align: left;
                   }}
                   th {{
                       background-color: {self.format_config['color_primary']};
                       color: white;
                   }}
                   tr:nth-child(even) {{
                       background-color: #f9f9f9;
                   }}
                   .image-container {{
                       text-align: center;
                       margin: 20px 0;
                   }}
                   .image-container img {{
                       max-width: 100%;
                       height: auto;
                   }}
                   .caption {{
                       font-style: italic;
                       text-align: center;
                       color: {self.format_config['color_text_light']};
                       margin-top: 5px;
                   }}
                   .header {{
                       text-align: center;
                       margin-bottom: 30px;
                   }}
                   .footer {{
                       text-align: center;
                       margin-top: 50px;
                       padding-top: 20px;
                       border-top: 1px solid #ddd;
                       color: {self.format_config['color_text_light']};
                   }}
                   .llm-insights {{
                       background-color: #f5f5f5;
                       border-left: 5px solid {self.format_config['color_secondary']};
                       padding: 15px;
                       margin: 20px 0;
                   }}
               </style>
           </head>
           <body>
               <div class="header">
                   <h1>{report_data['title']}</h1>
                   <p>{report_data['subtitle']}</p>
                   <p>Generated: {report_data['date']}</p>
               </div>
           """
           
           # Add each section
           for section_name, section in report_data['sections'].items():
               html_content += f"""
               <h2>{section['title']}</h2>
               """
               
               # Add section content
               for item in section['content']:
                   if item['type'] == 'paragraph':
                       html_content += f"""
                       <p>{item['text']}</p>
                       """
                   
                   elif item['type'] == 'subheading':
                       html_content += f"""
                       <h3>{item['text']}</h3>
                       """
                   
                   elif item['type'] == 'bullet_list':
                       html_content += """
                       <ul>
                       """
                       for bullet in item['items']:
                           html_content += f"""
                           <li>{bullet}</li>
                           """
                       html_content += """
                       </ul>
                       """
                   
                   elif item['type'] == 'numbered_list':
                       html_content += """
                       <ol>
                       """
                       for entry in item['items']:
                           html_content += f"""
                           <li>{entry}</li>
                           """
                       html_content += """
                       </ol>
                       """
                   
                   elif item['type'] == 'table':
                       html_content += """
                       <table>
                           <thead>
                               <tr>
                       """
                       for header in item['headers']:
                           html_content += f"""
                           <th>{header}</th>
                           """
                       html_content += """
                               </tr>
                           </thead>
                           <tbody>
                       """
                       for row in item['rows']:
                           html_content += """
                           <tr>
                           """
                           for cell in row:
                               html_content += f"""
                               <td>{cell}</td>
                               """
                           html_content += """
                           </tr>
                           """
                       html_content += """
                           </tbody>
                       </table>
                       """
                   
                   elif item['type'] == 'image':
                       if os.path.exists(item['path']):
                           # Convert the image path to a data URL for portability
                           with open(item['path'], 'rb') as img_file:
                               img_data = base64.b64encode(img_file.read()).decode('utf-8')
                               img_ext = os.path.splitext(item['path'])[1].lstrip('.')
                               data_url = f"data:image/{img_ext};base64,{img_data}"
                           
                           html_content += f"""
                           <div class="image-container">
                               <img src="{data_url}" alt="{section['title']} visualization">
                               <div class="caption">{item.get('caption', '')}</div>
                           </div>
                           """
               
               # Add LLM insights if available
               if 'llm_insights' in section:
                   html_content += f"""
                   <div class="llm-insights">
                       <p>{section['llm_insights']}</p>
                   </div>
                   """
               
               # Add variable insights if available
               if section_name == 'methodology' and 'variable_insights' in section:
                   html_content += f"""
                   <div class="llm-insights">
                       <h3>Variable Relationship Analysis</h3>
                       <p>{section['variable_insights']}</p>
                   </div>
                   """
               
               # Add pattern insights if available
               if section_name == 'vulnerability_rankings' and 'pattern_insights' in section:
                   html_content += f"""
                   <div class="llm-insights">
                       <h3>Vulnerability Pattern Analysis</h3>
                       <p>{section['pattern_insights']}</p>
                   </div>
                   """
           
           # Add footer
           html_content += f"""
               <div class="footer">
                   <p>Malaria Risk Analysis Report | Generated by MRPT | {report_data['date']}</p>
               </div>
           </body>
           </html>
           """
           
           if save_file:
               # Create unique filename
               timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
               filename = f"malaria_risk_report_{timestamp}.html"
               report_path = os.path.join(self.report_folder, filename)
               
               # Save the HTML file
               with open(report_path, 'w', encoding='utf-8') as f:
                   f.write(html_content)
               
               # Create report URL
               report_url = f"/download_report/{filename}"
               
               return report_path, report_url
           else:
               # Just return the content path (temp file)
               temp_path = os.path.join(self.report_folder, f"temp_report_{uuid.uuid4()}.html")
               with open(temp_path, 'w', encoding='utf-8') as f:
                   f.write(html_content)
               return temp_path, None
           
       except Exception as e:
           logger.error(f"Error generating HTML report: {str(e)}", exc_info=True)
           raise
   
   def _generate_docx_report(self, report_data):
       """
       Generate DOCX report
       
       Args:
           report_data: Data for the report
           
       Returns:
           tuple: (report_path, report_url)
       """
       try:
           # Create document
           doc = Document()
           
           # Set document properties
           doc.core_properties.title = report_data['title']
           doc.core_properties.subject = report_data['subtitle']
           doc.core_properties.created = datetime.now()
           
           # Add title
           title = doc.add_heading(report_data['title'], level=0)
           title.alignment = WD_ALIGN_PARAGRAPH.CENTER
           
           # Add subtitle
           subtitle = doc.add_paragraph(report_data['subtitle'])
           subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
           subtitle.style = 'Subtitle'
           
           # Add date
           date_paragraph = doc.add_paragraph(f"Generated: {report_data['date']}")
           date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
           
           # Add page break after title page
           doc.add_page_break()
           
           # Add each section
           for section_name, section in report_data['sections'].items():
               # Section heading
               doc.add_heading(section['title'], level=1)
               
               # Add section content
               for item in section['content']:
                   if item['type'] == 'paragraph':
                       doc.add_paragraph(item['text'])
                   
                   elif item['type'] == 'subheading':
                       doc.add_heading(item['text'], level=2)
                   
                   elif item['type'] == 'bullet_list':
                       for bullet in item['items']:
                           paragraph = doc.add_paragraph(bullet)
                           paragraph.style = 'List Bullet'
                   
                   elif item['type'] == 'numbered_list':
                       for entry in item['items']:
                           paragraph = doc.add_paragraph(entry)
                           paragraph.style = 'List Number'
                   
                   elif item['type'] == 'table':
                       # Create table
                       table = doc.add_table(rows=1, cols=len(item['headers']))
                       table.style = 'Table Grid'
                       
                       # Add headers
                       header_cells = table.rows[0].cells
                       for i, header in enumerate(item['headers']):
                           header_cells[i].text = header
                           # Make headers bold
                           for paragraph in header_cells[i].paragraphs:
                               for run in paragraph.runs:
                                   run.bold = True
                       
                       # Add data rows
                       for row_data in item['rows']:
                           row_cells = table.add_row().cells
                           for i, cell_data in enumerate(row_data):
                               row_cells[i].text = cell_data
                       
                       # Add space after table
                       doc.add_paragraph()
                   
                   elif item['type'] == 'image':
                       if os.path.exists(item['path']):
                           # Add image with caption
                           doc.add_picture(item['path'], width=Inches(6))
                           if 'caption' in item:
                               caption = doc.add_paragraph(item['caption'])
                               caption.style = 'Caption'
                               caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
               
               # Add LLM insights if available
               if 'llm_insights' in section:
                   # Add styled paragraph for LLM insights
                   doc.add_heading('Analysis Insights', level=3)
                   insights = doc.add_paragraph(section['llm_insights'])
                   
                   # Style the paragraph as a callout
                   for run in insights.runs:
                       run.font.italic = True
               
               # Add variable insights if available
               if section_name == 'methodology' and 'variable_insights' in section:
                   doc.add_heading('Variable Relationship Analysis', level=3)
                   insights = doc.add_paragraph(section['variable_insights'])
                   for run in insights.runs:
                       run.font.italic = True
               
               # Add pattern insights if available
               if section_name == 'vulnerability_rankings' and 'pattern_insights' in section:
                   doc.add_heading('Vulnerability Pattern Analysis', level=3)
                   insights = doc.add_paragraph(section['pattern_insights'])
                   for run in insights.runs:
                       run.font.italic = True
               
               # Add page break after each section except the last one
               if section_name != list(report_data['sections'].keys())[-1]:
                   doc.add_page_break()
           
           # Add footer
           section = doc.sections[-1]
           footer = section.footer
           footer_text = footer.paragraphs[0]
           footer_text.text = f"Malaria Risk Analysis Report | Generated by MRPT | {report_data['date']}"
           footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
           
           # Create unique filename
           timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
           filename = f"malaria_risk_report_{timestamp}.docx"
           report_path = os.path.join(self.report_folder, filename)
           
           # Save the document
           doc.save(report_path)
           
           # Create report URL
           report_url = f"/download_report/{filename}"
           
           return report_path, report_url
           
       except Exception as e:
           logger.error(f"Error generating DOCX report: {str(e)}", exc_info=True)
           raise
   
   def _generate_markdown_report(self, report_data):
       """
       Generate Markdown report
       
       Args:
           report_data: Data for the report
           
       Returns:
           tuple: (report_path, report_url)
       """
       try:
           # Start building markdown content
           markdown_content = f"# {report_data['title']}\n\n"
           markdown_content += f"_{report_data['subtitle']}_\n\n"
           markdown_content += f"Generated: {report_data['date']}\n\n"
           
           # Add horizontal rule
           markdown_content += "---\n\n"
           
           # Add each section
           for section_name, section in report_data['sections'].items():
               # Section heading
               markdown_content += f"## {section['title']}\n\n"
               
               # Add section content
               for item in section['content']:
                   if item['type'] == 'paragraph':
                       markdown_content += f"{item['text']}\n\n"
                   
                   elif item['type'] == 'subheading':
                       markdown_content += f"### {item['text']}\n\n"
                   
                   elif item['type'] == 'bullet_list':
                       for bullet in item['items']:
                           markdown_content += f"* {bullet}\n"
                       markdown_content += "\n"
                   
                   elif item['type'] == 'numbered_list':
                       for i, entry in enumerate(item['items']):
                           markdown_content += f"{i+1}. {entry}\n"
                       markdown_content += "\n"
                   
                   elif item['type'] == 'table':
                       # Add table headers
                       markdown_content += "| " + " | ".join(item['headers']) + " |\n"
                       # Add separator row
                       markdown_content += "| " + " | ".join(["---"] * len(item['headers'])) + " |\n"
                       # Add data rows
                       for row in item['rows']:
                           markdown_content += "| " + " | ".join(row) + " |\n"
                       markdown_content += "\n"
                   
                   elif item['type'] == 'image':
                       if os.path.exists(item['path']):
                           # Images in markdown will be references to relative paths
                           # For a downloadable report, we'll need to include the images separately
                           # For now, just reference the path
                           image_filename = os.path.basename(item['path'])
                           markdown_content += f"![{section['title']} visualization]({image_filename})\n\n"
                           if 'caption' in item:
                               markdown_content += f"*{item['caption']}*\n\n"
               
               # Add LLM insights if available
               if 'llm_insights' in section:
                   markdown_content += f"> **Analysis Insights:**\n> \n> {section['llm_insights']}\n\n"
               
               # Add variable insights if available
               if section_name == 'methodology' and 'variable_insights' in section:
                   markdown_content += f"> **Variable Relationship Analysis:**\n> \n> {section['variable_insights']}\n\n"
               
               # Add pattern insights if available
               if section_name == 'vulnerability_rankings' and 'pattern_insights' in section:
                   markdown_content += f"> **Vulnerability Pattern Analysis:**\n> \n> {section['pattern_insights']}\n\n"
               
               # Add horizontal rule between sections
               markdown_content += "---\n\n"
           
           # Add footer
           markdown_content += f"*Malaria Risk Analysis Report | Generated by MRPT | {report_data['date']}*"
           
           # Create unique filename
           timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
           filename = f"malaria_risk_report_{timestamp}.md"
           report_path = os.path.join(self.report_folder, filename)
           
           # Save the markdown file
           with open(report_path, 'w', encoding='utf-8') as f:
               f.write(markdown_content)
           
           # Create report URL
           report_url = f"/download_report/{filename}"
           
           return report_path, report_url
           
       except Exception as e:
           logger.error(f"Error generating Markdown report: {str(e)}", exc_info=True)
           raise
   
   def _log_report_generation(self, format_type, report_path):
       """
       Log report generation event
       
       Args:
           format_type: Report format
           report_path: Path to the generated report
       """
       try:
           # Get interaction logger if available
           interaction_logger = current_app.config.get('INTERACTION_LOGGER')
           if interaction_logger and self.session_id:
               # Get file size
               file_size = os.path.getsize(report_path) if os.path.exists(report_path) else 0
               
               # Log the event
               interaction_logger.log_analysis_event(
                   self.session_id,
                   'report_generation',
                   {
                       'format': format_type,
                       'file_path': report_path,
                       'file_size': file_size,
                       'timestamp': datetime.now().isoformat(),
                       'detail_level': self.metadata.get('detail_level', 'standard') if hasattr(self, 'metadata') else 'standard'
                   },
                   True  # Success
               )
       except Exception as e:
           logger.error(f"Error logging report generation: {str(e)}")


def generate_report(data_handler, format_type='pdf', custom_sections=None, detail_level='standard'):
   """
   User-friendly function to generate a report from the data handler
   
   Args:
       data_handler: DataHandler instance
       format_type: Type of report to generate (pdf, html, docx, md)
       custom_sections: Optional list of custom sections to include
       detail_level: Level of detail (basic, standard, technical)
       
   Returns:
       dict: Status and report information
   """
   try:
       # Create report generator
       report_gen = ReportGenerator(data_handler)
       
       # Generate the report
       result = report_gen.generate_report(format_type, custom_sections, detail_level)
       
       return result
   except Exception as e:
       logger.error(f"Error generating report: {str(e)}", exc_info=True)
       error_info = log_error_details(e, 
                                    session_id=os.path.basename(data_handler.session_folder), 
                                    context={'format': format_type})
       return {
           'status': 'error',
           'message': f'Error generating report: {str(e)}',
           'error_details': error_info
       }